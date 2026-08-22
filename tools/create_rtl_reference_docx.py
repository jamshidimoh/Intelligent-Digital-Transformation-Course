from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_bidi(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)


def set_rtl_style(style, size, bold=False):
    style.font.name = 'Noto Naskh Arabic'
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Noto Naskh Arabic')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Noto Naskh Arabic')
    style._element.rPr.rFonts.set(qn('w:cs'), 'Noto Naskh Arabic')
    style.font.size = Pt(size)
    style.font.bold = bold
    ppr = style._element.get_or_add_pPr()
    bidi = ppr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        ppr.append(bidi)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(2.2)

for name, size, bold in [('Normal', 12.5, False), ('Title', 22, True), ('Heading 1', 17, True), ('Heading 2', 15, True), ('Heading 3', 13.5, True)]:
    set_rtl_style(doc.styles[name], size, bold)
    doc.styles[name].paragraph_format.space_after = Pt(7)
    doc.styles[name].paragraph_format.line_spacing = 1.35

doc.add_paragraph('قالب مرجع RTL', style='Title')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_bidi(doc.paragraphs[-1])
path = 'tools/rtl-reference.docx'
doc.save(path)
print(path)
