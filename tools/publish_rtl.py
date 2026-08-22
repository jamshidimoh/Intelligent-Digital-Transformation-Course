from __future__ import annotations

from pathlib import Path
import html
import re
import subprocess

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / 'PUBLISH' / 'CHAPTERS'


def set_bidi(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_font(run, size=12.5, bold=False):
    run.font.name = 'Noto Naskh Arabic'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Noto Naskh Arabic')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Noto Naskh Arabic')
    run._element.rPr.rFonts.set(qn('w:cs'), 'Noto Naskh Arabic')
    run.font.size = Pt(size)
    run.bold = bold


def isolate_latin(text: str) -> str:
    text = text.replace('\u200e', '').replace('\u200f', '')
    return re.sub(r'([A-Za-z][A-Za-z0-9._+/#:-]*(?:\s+[A-Za-z][A-Za-z0-9._+/#:-]*)*)', r'\u200e\1\u200e', text)


def read_chapter(chapter_dir: Path) -> str:
    parts = [chapter_dir / 'chapter.md']
    parts += sorted((chapter_dir / 'sections').glob('*.md'))
    refs = chapter_dir / 'references.md'
    if refs.exists():
        parts.append(refs)
    return '\n\n'.join(p.read_text(encoding='utf-8') for p in parts)


def add_docx_content(doc: Document, text: str, chapter_dir: Path):
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith('# '):
            p = doc.add_paragraph(style='Title')
            set_bidi(p); r = p.add_run(isolate_latin(line[2:])); set_font(r, 20, True)
        elif line.startswith('## '):
            p = doc.add_paragraph(style='Heading 1')
            set_bidi(p); r = p.add_run(isolate_latin(line[3:])); set_font(r, 16, True)
        elif line.startswith('### '):
            p = doc.add_paragraph(style='Heading 2')
            set_bidi(p); r = p.add_run(isolate_latin(line[4:])); set_font(r, 14, True)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            set_bidi(p); r = p.add_run(isolate_latin(line[2:])); set_font(r)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            set_bidi(p); r = p.add_run(isolate_latin(re.sub(r'^\d+\.\s*', '', line))); set_font(r)
        elif line.startswith('!['):
            # Images are handled by the PDF/HTML path; Word keeps the text caption.
            m = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if m:
                p = doc.add_paragraph()
                set_bidi(p); r = p.add_run('شکل: ' + m.group(1)); set_font(r, 11, False)
        elif line == '---':
            p = doc.add_paragraph(); set_bidi(p)
        else:
            p = doc.add_paragraph()
            set_bidi(p)
            p.paragraph_format.line_spacing = 1.35
            r = p.add_run(isolate_latin(line)); set_font(r)


def publish_docx(chapter_dir: Path, text: str, out_dir: Path):
    doc = Document()
    for name, size, bold in [('Normal', 12.5, False), ('Title', 20, True), ('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 13, True)]:
        st = doc.styles[name]
        st.font.name = 'Noto Naskh Arabic'
        st._element.rPr.rFonts.set(qn('w:ascii'), 'Noto Naskh Arabic')
        st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Noto Naskh Arabic')
        st._element.rPr.rFonts.set(qn('w:cs'), 'Noto Naskh Arabic')
        st.font.size = Pt(size); st.font.bold = bold
    add_docx_content(doc, text, chapter_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / 'chapter.docx'
    doc.save(path)
    return path


def publish_pdf(chapter_dir: Path, text: str, out_dir: Path):
    # HTML/CSS is used for PDF because it supports explicit RTL layout reliably.
    body=[]
    for raw in text.splitlines():
        line=raw.strip()
        if not line: continue
        if line.startswith('# '): body.append(f'<h1>{html.escape(isolate_latin(line[2:]))}</h1>')
        elif line.startswith('## '): body.append(f'<h2>{html.escape(isolate_latin(line[3:]))}</h2>')
        elif line.startswith('### '): body.append(f'<h3>{html.escape(isolate_latin(line[4:]))}</h3>')
        elif line.startswith('- '): body.append(f'<p class="bullet">• {html.escape(isolate_latin(line[2:]))}</p>')
        elif re.match(r'^\d+\.\s', line): body.append(f'<p class="bullet">{html.escape(isolate_latin(line))}</p>')
        elif line == '---': body.append('<hr>')
        elif line.startswith('!['):
            m=re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if m:
                src=(chapter_dir / m.group(2)).resolve().as_uri()
                body.append(f'<figure><img src="{src}"><figcaption>{html.escape(m.group(1))}</figcaption></figure>')
        else: body.append(f'<p>{html.escape(isolate_latin(line))}</p>')
    css=(ROOT/'tools'/'rtl.css').read_text(encoding='utf-8')
    html_doc=f'<!doctype html><html lang="fa" dir="rtl"><head><meta charset="utf-8"><style>{css}</style></head><body>{"".join(body)}</body></html>'
    html_path=out_dir/'chapter.html'; out_dir.mkdir(parents=True, exist_ok=True); html_path.write_text(html_doc, encoding='utf-8')
    pdf_path=out_dir/'chapter.pdf'
    try:
        from weasyprint import HTML
        HTML(filename=str(html_path)).write_pdf(str(pdf_path))
    except Exception as exc:
        print(f'WeasyPrint unavailable/failed: {exc}')
        return None
    return pdf_path


def main():
    chapter_dir=ROOT/'BOOK'/'01-foundations'
    text=read_chapter(chapter_dir)
    out=ROOT/'PUBLISH'/'CHAPTERS'/'01-foundations'
    docx=publish_docx(chapter_dir, text, out)
    pdf=publish_pdf(chapter_dir, text, out)
    print(docx)
    print(pdf)

if __name__ == '__main__':
    main()
