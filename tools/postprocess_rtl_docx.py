from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import sys
import unicodedata

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PERSIAN_FONT = "Noto Naskh Arabic"
LATIN_FONT = "Noto Sans"
ACCENT = RGBColor(31, 61, 82)
FA_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]")
LATIN_RE = re.compile(r"[A-Za-z]")
NUMBER_RE = re.compile(r"\d")


def set_bool(parent, tag: str, enabled: bool = True) -> None:
    node = parent.find(qn(tag))
    if enabled and node is None:
        parent.append(OxmlElement(tag))
    elif not enabled and node is not None:
        parent.remove(node)


def set_lang(rpr, val: str) -> None:
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), val)


def set_run_font(rpr, rtl: bool) -> None:
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), PERSIAN_FONT if rtl else LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), LATIN_FONT)


def set_run_direction(run_el, rtl: bool) -> None:
    rpr = run_el.get_or_add_rPr()
    set_bool(rpr, "w:rtl", rtl)
    set_lang(rpr, "fa-IR" if rtl else "en-US")
    set_run_font(rpr, rtl)


def get_text(run_el) -> str:
    return "".join((node.text or "") for node in run_el.findall(".//" + qn("w:t")))


def classify_char(ch: str) -> str:
    if FA_RE.match(ch):
        return "fa"
    if LATIN_RE.match(ch) or NUMBER_RE.match(ch):
        return "en"
    if unicodedata.category(ch).startswith("L"):
        return "fa"
    return "neutral"


def tokenize_mixed(text: str) -> list[tuple[str, str]]:
    chunks: list[tuple[str, str]] = []
    current: list[str] = []
    current_kind: str | None = None
    for ch in text:
        kind = classify_char(ch)
        if kind == "neutral":
            if current_kind is None:
                current_kind = "fa"
            current.append(ch)
            continue
        if current_kind is None:
            current_kind = kind
        elif kind != current_kind:
            chunks.append((current_kind, "".join(current)))
            current = []
            current_kind = kind
        current.append(ch)
    if current:
        chunks.append((current_kind or "fa", "".join(current)))
    return [(k, t) for k, t in chunks if t]


def make_run_clone(run_el, text: str, rtl: bool):
    clone = deepcopy(run_el)
    t_nodes = clone.findall(".//" + qn("w:t"))
    if not t_nodes:
        return clone
    for node in t_nodes:
        node.text = ""
    t_nodes[0].text = text
    set_run_direction(clone, rtl)
    return clone


def replace_run_with_script_chunks(run) -> None:
    run_el = run._r
    text = get_text(run_el)
    if not text:
        return
    parent = run_el.getparent()
    if parent is None:
        return
    chunks = tokenize_mixed(text)
    meaningful = [(k, t) for k, t in chunks if any(c.isalnum() for c in t)]
    if len(meaningful) <= 1:
        set_run_direction(run_el, meaningful[0][0] == "fa" if meaningful else True)
        return
    index = parent.index(run_el)
    for kind, piece in chunks:
        parent.insert(index, make_run_clone(run_el, piece, kind == "fa"))
        index += 1
    parent.remove(run_el)


def set_paragraph_rtl(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    set_bool(ppr, "w:bidi", True)
    set_bool(ppr, "w:widowControl", True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def paragraph_is_code_like(paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    text = paragraph.text or ""
    return "code" in style_name or text.strip().startswith(("```", "$ ", ">> "))


def tune_styles(doc: Document) -> None:
    styles = doc.styles
    if "Normal" in styles:
        s = styles["Normal"]
        s.font.name = PERSIAN_FONT
        s.font.size = Pt(12.5)
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.34
        s.paragraph_format.first_line_indent = Cm(0.42)
        set_bool(s._element.get_or_add_pPr(), "w:bidi", False)

    heading_specs = {
        "Heading 1": (17, 18, 10, 1.15),
        "Heading 2": (14.5, 13, 7, 1.15),
        "Heading 3": (13.3, 10, 5, 1.15),
        "Heading 4": (12.5, 8, 4, 1.15),
    }
    for name, (size, before, after, line) in heading_specs.items():
        if name in styles:
            s = styles[name]
            s.font.name = PERSIAN_FONT
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = ACCENT
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(after)
            s.paragraph_format.line_spacing = line
            s.paragraph_format.first_line_indent = Cm(0)
            ppr = s._element.get_or_add_pPr()
            set_bool(ppr, "w:bidi", False)
            set_bool(ppr, "w:keepNext", True)
            set_bool(ppr, "w:widowControl", True)

    for name in ("List Bullet", "List Number", "List Bullet 2", "List Number 2", "Quote", "Intense Quote", "Caption"):
        if name in styles:
            s = styles[name]
            s.font.name = PERSIAN_FONT
            s.paragraph_format.first_line_indent = Cm(0)
            set_bool(s._element.get_or_add_pPr(), "w:bidi", False)
            set_bool(s._element.get_or_add_pPr(), "w:widowControl", True)


def tune_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.35)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.55)
        section.right_margin = Cm(2.35)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.9)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def paragraph_is_caption(paragraph) -> bool:
    return paragraph.style.name == "Caption" if paragraph.style else False


def configure_paragraph(paragraph) -> None:
    if paragraph_is_code_like(paragraph):
        ppr = paragraph._p.get_or_add_pPr()
        set_bool(ppr, "w:bidi", False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        for run in list(paragraph.runs):
            set_run_direction(run._r, False)
        return

    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0.42)

    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        set_bool(paragraph._p.get_or_add_pPr(), "w:keepNext", True)
    elif paragraph_has_drawing(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(5)
    elif paragraph_is_caption(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(8)
    elif style_name.startswith("List "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.4)
        paragraph.paragraph_format.right_indent = Cm(0.2)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for run in list(paragraph.runs):
        replace_run_with_script_chunks(run)

    ppr = paragraph._p.get_or_add_pPr()
    set_bool(ppr, "w:widowControl", True)


def configure_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    set_bool(tbl_pr, "w:bidiVisual", True)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            rotated = tc_pr.find(qn("w:textDirection"))
            if rotated is not None:
                tc_pr.remove(rotated)
            for paragraph in cell.paragraphs:
                set_paragraph_rtl(paragraph)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in list(paragraph.runs):
                    replace_run_with_script_chunks(run)
                if row_index == 0:
                    for run in paragraph.runs:
                        run.bold = True


def configure_header_footer(paragraph) -> None:
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    for run in list(paragraph.runs):
        replace_run_with_script_chunks(run)


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: postprocess_rtl_docx.py <docx>", file=sys.stderr)
        return 2
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"Missing DOCX: {path}", file=sys.stderr)
        return 1

    doc = Document(path)
    tune_sections(doc)
    tune_styles(doc)
    for paragraph in doc.paragraphs:
        configure_paragraph(paragraph)
    for table in doc.tables:
        configure_table(table)
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            configure_header_footer(paragraph)

    doc.save(path)
    print(f"Post-processed whole-document RTL academic DOCX: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
