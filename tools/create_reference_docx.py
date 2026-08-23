from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "PUBLISH" / "reference-template.docx"

PERSIAN_FONT = "Noto Naskh Arabic"
LATIN_FONT = "Noto Sans"
ACCENT = RGBColor(31, 61, 82)


def add_bool_property(parent, tag: str) -> None:
    if parent.find(qn(tag)) is None:
        parent.append(OxmlElement(tag))


def set_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    add_bool_property(p_pr, "w:bidi")


def configure_style(style, size: float, bold: bool = False, color: RGBColor | None = None,
                    before: float = 0, after: float = 7, line: float = 1.28) -> None:
    style.font.name = PERSIAN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = color
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), PERSIAN_FONT)
    rfonts.set(qn("w:eastAsia"), LATIN_FONT)
    ppr = style._element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(round(before * 20)))
    spacing.set(qn("w:after"), str(round(after * 20)))
    spacing.set(qn("w:line"), str(round(line * 240)))
    spacing.set(qn("w:lineRule"), "auto")
    add_bool_property(ppr, "w:bidi")


def add_keep_with_next(style) -> None:
    ppr = style._element.get_or_add_pPr()
    add_bool_property(ppr, "w:keepNext")


def add_widow_control(style) -> None:
    ppr = style._element.get_or_add_pPr()
    add_bool_property(ppr, "w:widowControl")


def configure_cell_style(style) -> None:
    style.font.name = PERSIAN_FONT
    style.font.size = Pt(10.5)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), PERSIAN_FONT)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    configure_style(styles["Normal"], 12.2, False, None, 0, 7, 1.28)
    styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["Normal"].paragraph_format.first_line_indent = Cm(0.45)
    add_widow_control(styles["Normal"])

    configure_style(styles["Title"], 21, True, ACCENT, 0, 10, 1.15)
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    styles["Title"].paragraph_format.first_line_indent = Cm(0)
    add_keep_with_next(styles["Title"])

    if "Subtitle" in styles:
        configure_style(styles["Subtitle"], 14, False, ACCENT, 0, 16, 1.2)
        styles["Subtitle"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        styles["Subtitle"].paragraph_format.first_line_indent = Cm(0)
        add_keep_with_next(styles["Subtitle"])

    configure_style(styles["Heading 1"], 17, True, ACCENT, 16, 8, 1.15)
    configure_style(styles["Heading 2"], 14.5, True, ACCENT, 12, 6, 1.15)
    configure_style(styles["Heading 3"], 13.2, True, ACCENT, 10, 5, 1.15)
    configure_style(styles["Heading 4"], 12.3, True, ACCENT, 8, 4, 1.15)
    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        styles[name].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        styles[name].paragraph_format.first_line_indent = Cm(0)
        add_keep_with_next(styles[name])

    if "Caption" in styles:
        configure_style(styles["Caption"], 10.5, False, RGBColor(70, 70, 70), 4, 8, 1.15)
        styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styles["Caption"].paragraph_format.first_line_indent = Cm(0)
        add_keep_with_next(styles["Caption"])

    if "Table Contents" in styles:
        configure_cell_style(styles["Table Contents"])
    if "Table Heading" in styles:
        configure_cell_style(styles["Table Heading"])
        styles["Table Heading"].font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(header)
    r = header.add_run("تحول دیجیتال هوشمند | معماری، چارچوب‌ها و پیاده‌سازی")
    r.font.name = PERSIAN_FONT
    r.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(footer)
    r = footer.add_run("صفحه ")
    r.font.name = PERSIAN_FONT
    r.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    # Ensure template styles exist; content is supplied later by Pandoc.
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
